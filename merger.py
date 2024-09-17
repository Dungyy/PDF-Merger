import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QListWidget, QPushButton, QVBoxLayout, QWidget, 
    QFileDialog, QMessageBox, QHBoxLayout, QProgressBar, QLabel, QGridLayout
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
from lxml import etree
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import tempfile
import comtypes.client

class FileProcessorThread(QThread):
    progress = pyqtSignal(int)
    error = pyqtSignal(str)

    def __init__(self, function, *args, **kwargs):
        super().__init__()
        self.function = function
        self.args = args
        self.kwargs = kwargs

    def run(self):
        try:
            self.function(*self.args, **self.kwargs)
        except Exception as e:
            self.error.emit(str(e))

class FileMerger(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Dungerz File Merger: PDFs, DOCX, Images & More")
        self.setGeometry(100, 100, 800, 550)
        self.setAcceptDrops(True)  # Enable drag and drop for the entire window
        self.dark_mode = True  # Start in dark mode by default
        self.initUI()

    def initUI(self):
        layout = QGridLayout()

        # Title and Description Labels
        self.title_label = QLabel("Dungerz File Merger")
        self.title_label.setAlignment(Qt.AlignCenter)
        self.title_label.setStyleSheet("font-size: 30px; font-weight: bold; margin-bottom: 5px;")

        self.description_label = QLabel(
            "Easily merge PDFs, Word documents, images, and XML files into a single file. "
            "Drag and drop your files, or use the buttons below to add them."
        )
        self.description_label.setAlignment(Qt.AlignCenter)
        self.description_label.setWordWrap(True)

        # List widget to display files
        self.listWidget = QListWidget()
        self.listWidget.setDragEnabled(True)
        self.listWidget.setSelectionMode(QListWidget.MultiSelection)
        self.listWidget.setDragDropMode(QListWidget.DragDropMode.InternalMove)

        # Progress Bar
        self.progressBar = QProgressBar()
        self.progressBar.setTextVisible(True)
        self.progressBar.setMaximum(100)
        self.progressBar.setValue(0)

        # Buttons Layout for bottom
        button_layout = QHBoxLayout()
        self.addButton = self.create_button("Add Files", self.open_file_explorer)
        self.mergeButton = self.create_button("Merge and Save", self.merge_files)
        self.clearButton = self.create_button("Clear List", self.clear_list)
        self.removeSelectedButton = self.create_button("Remove Selected", self.remove_selected_items)

        button_layout.addWidget(self.addButton)
        button_layout.addWidget(self.mergeButton)
        button_layout.addWidget(self.removeSelectedButton)
        button_layout.addWidget(self.clearButton)

        # Dark/Light Mode Toggle at the top right
        self.switch_mode_button = self.create_button("Toggle Mode", self.switch_mode)
        self.switch_mode_button.setFixedWidth(120)

        # Top Right corner layout for mode switch
        top_layout = QHBoxLayout()
        top_layout.addStretch()
        top_layout.addWidget(self.switch_mode_button)

        # Set Layouts
        layout.addLayout(top_layout, 0, 2)
        layout.addWidget(self.title_label, 0, 0, 1, 2)
        layout.addWidget(self.description_label, 1, 0, 1, 3)
        layout.addWidget(self.listWidget, 2, 0, 1, 3)
        layout.addWidget(self.progressBar, 3, 0, 1, 3)
        layout.addLayout(button_layout, 4, 0, 1, 3)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.setup_styles()

    def create_button(self, text, callback):
        button = QPushButton(text)
        button.clicked.connect(callback)
        return button

    def setup_styles(self):
        if self.dark_mode:
            self.setStyleSheet("""
                QMainWindow {
                    background-color: #2e2e2e;
                }
                QLabel {
                    color: #ffffff;
                }
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    font-size: 16px;
                    padding: 10px;
                    border-radius: 8px;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
                QListWidget {
                    background-color: #1e1e1e;
                    color: white;
                    border: 2px dashed #cccccc;
                }
                QProgressBar {
                    background-color: #1e1e1e;
                    color: white;
                    border: 1px solid #cccccc;
                    height: 20px;
                }
            """)
        else:
            self.setStyleSheet("""
                QMainWindow {
                    background-color: #ffffff;
                }
                QLabel {
                    color: #000000;
                }
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    font-size: 16px;
                    padding: 10px;
                    border-radius: 8px;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
                QListWidget {
                    background-color: #f0f0f0;
                    color: black;
                    border: 2px dashed #cccccc;
                }
                QProgressBar {
                    background-color: #f0f0f0;
                    color: black;
                    border: 1px solid #000000;
                    height: 20px;
                }
            """)

    def switch_mode(self):
        self.dark_mode = not self.dark_mode
        self.setup_styles()

    # Drag and Drop Events
    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        files = [u.toLocalFile() for u in event.mimeData().urls()]
        self.add_files(files)

    def open_file_explorer(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "Select Files", "", 
            "All Files (*);;PDF Files (*.pdf);;Word Files (*.docx);;Image Files (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;XML Files (*.xml)"
        )
        self.add_files(files)

    def add_files(self, files):
        for file in files:
            if file.lower().endswith(('.pdf', '.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.docx', '.xml')):
                items = self.listWidget.findItems(file, Qt.MatchExactly)
                if not items:
                    self.listWidget.addItem(file)
            else:
                QMessageBox.warning(self, "Unsupported File", f"The file {os.path.basename(file)} is not supported.")

    # Merge the selected files and export to desired format (PDF, DOCX, etc.)
    def merge_files(self):
        if self.listWidget.count() == 0:
            QMessageBox.information(self, "No Files Selected", "Please add files to merge.")
            return
        
        export_formats = "PDF Files (*.pdf);;Word Files (*.docx);;All Files (*)"
        output_filename, selected_format = QFileDialog.getSaveFileName(self, "Save Merged File", "", export_formats)
        if not output_filename:
            return

        _, file_extension = os.path.splitext(output_filename)
        if file_extension.lower() == '.pdf':
            self.process_files(self.export_to_pdf, output_filename)
        elif file_extension.lower() == '.docx':
            self.process_files(self.export_to_docx, output_filename)
        else:
            QMessageBox.warning(self, "Unsupported Format", f"The format {file_extension} is not supported.")

    def process_files(self, export_function, output_filename):
        self.thread = FileProcessorThread(export_function, output_filename)
        self.thread.progress.connect(self.update_progress)
        self.thread.error.connect(self.show_error_message)
        self.thread.finished.connect(self.on_merge_finished)
        self.thread.start()
        self.mergeButton.setEnabled(False)

    def update_progress(self, value):
        self.progressBar.setValue(value)

    def show_error_message(self, message):
        QMessageBox.warning(self, "Merge Failed", f"An error occurred during the merge process: {message}")

    def on_merge_finished(self):
        self.progressBar.setValue(0)
        self.mergeButton.setEnabled(True)
        QMessageBox.information(self, "Merge Successful", f"Merged file saved successfully.")
        self.listWidget.clear()

    def export_to_pdf(self, output_filename):
        pdf_writer = PdfWriter()
        num_files = self.listWidget.count()
        for index in range(num_files):
            filepath = self.listWidget.item(index).text()
            self.add_file_to_pdf(pdf_writer, filepath)
            self.thread.progress.emit(int((index + 1) / num_files * 100))

        with open(output_filename, 'wb') as out:
            pdf_writer.write(out)

    def add_file_to_pdf(self, pdf_writer, filepath):
        _, file_extension = os.path.splitext(filepath)
        
        if file_extension.lower() == '.pdf':
            self.add_pdf_to_pdf(pdf_writer, filepath)
        elif file_extension.lower() in ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'):
            self.add_image_to_pdf(pdf_writer, filepath)
        elif file_extension.lower() == '.docx':
            self.add_docx_to_pdf(pdf_writer, filepath)
        elif file_extension.lower() == '.xml':
            self.add_xml_to_pdf(pdf_writer, filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def add_pdf_to_pdf(self, pdf_writer, filepath):
        pdf_reader = PdfReader(filepath)
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)

    def add_image_to_pdf(self, pdf_writer, filepath):
        image = Image.open(filepath)
        pdf_bytes = BytesIO()
        image.convert("RGB").save(pdf_bytes, format='PDF')
        pdf_bytes.seek(0)
        temp_reader = PdfReader(pdf_bytes)
        for page in temp_reader.pages:
            pdf_writer.add_page(page)

    def add_docx_to_pdf(self, pdf_writer, filepath):
        word = comtypes.client.CreateObject('Word.Application')
        doc = word.Documents.Open(filepath)
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf.close()
        doc.SaveAs(temp_pdf.name, FileFormat=17)
        doc.Close()
        word.Quit()

        temp_reader = PdfReader(temp_pdf.name)
        for page in temp_reader.pages:
            pdf_writer.add_page(page)

        os.unlink(temp_pdf.name)

    def add_xml_to_pdf(self, pdf_writer, filepath):
        tree = etree.parse(filepath)
        root = tree.getroot()

        pdf_bytes = BytesIO()
        c = canvas.Canvas(pdf_bytes, pagesize=letter)
        width, height = letter
        y = height - 50

        def add_element(element, indent=0):
            nonlocal y
            if y < 50:
                c.showPage()
                y = height - 50
            tag = element.tag
            text = element.text.strip() if element.text else ""
            c.drawString(50 + indent * 20, y, f"{tag}: {text}")
            y -= 20
            for child in element:
                add_element(child, indent + 1)

        add_element(root)
        c.save()

        pdf_bytes.seek(0)
        temp_reader = PdfReader(pdf_bytes)
        for page in temp_reader.pages:
            pdf_writer.add_page(page)

    def export_to_docx(self, output_filename):
        doc = Document()
        num_files = self.listWidget.count()
        for index in range(num_files):
            filepath = self.listWidget.item(index).text()
            self.add_file_to_docx(doc, filepath)
            self.thread.progress.emit(int((index + 1) / num_files * 100))

        doc.save(output_filename)

    def add_file_to_docx(self, doc, filepath):
        _, file_extension = os.path.splitext(filepath)
        
        if file_extension.lower() == '.pdf':
            self.add_pdf_to_docx(doc, filepath)
        elif file_extension.lower() in ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'):
            self.add_image_to_docx(doc, filepath)
        elif file_extension.lower() == '.docx':
            self.add_docx_to_docx(doc, filepath)
        elif file_extension.lower() == '.xml':
            self.add_xml_to_docx(doc, filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def add_pdf_to_docx(self, doc, filepath):
        pdf_reader = PdfReader(filepath)
        for page in pdf_reader.pages:
            doc.add_paragraph(page.extract_text())

    def add_image_to_docx(self, doc, filepath):
        doc.add_picture(filepath, width=Inches(6))

    def add_docx_to_docx(self, doc, filepath):
        src_doc = Document(filepath)
        for element in src_doc.element.body:
            doc.element.body.append(element)

    def add_xml_to_docx(self, doc, filepath):
        tree = etree.parse(filepath)
        root = tree.getroot()

        def add_element(element, indent=0):
            tag = element.tag
            text = element.text.strip() if element.text else ""
            doc.add_paragraph(f"{'  ' * indent}{tag}: {text}")
            for child in element:
                add_element(child, indent + 1)

        add_element(root)

    def clear_list(self):
        self.listWidget.clear()

    def remove_selected_items(self):
        for selectedItem in self.listWidget.selectedItems():
            self.listWidget.takeItem(self.listWidget.row(selectedItem))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = FileMerger()
    window.show()
    sys.exit(app.exec_())
